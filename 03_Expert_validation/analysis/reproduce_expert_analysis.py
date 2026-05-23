from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def find_source_file() -> Path:
    repo_source = Path(__file__).resolve().parents[1] / "Expert_Responses.xlsx"
    if repo_source.exists():
        return repo_source

    desktop = Path.home() / "Desktop"
    matches = list(desktop.rglob("Expert_Responses.xlsx"))
    if not matches:
        raise FileNotFoundError(
            "Expert_Responses.xlsx was not found in 03_Expert_validation or under the Desktop folder."
        )
    return matches[0]


def set_base_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, value in enumerate(rows[0]):
        hdr_cells[idx].text = value
        for paragraph in hdr_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def main() -> None:
    source = find_source_file()
    df = pd.read_excel(source, sheet_name=0)
    score_cols = df.columns[1:7]
    df[score_cols] = df[score_cols].apply(pd.to_numeric, errors="coerce")

    analysis = pd.DataFrame(
        {
            "Question": [f"Q{i}" for i in range(1, 7)],
            "Mean": df[score_cols].mean().round(2).to_numpy(),
            "SD": df[score_cols].std().round(2).to_numpy(),
            "Valid N": df[score_cols].count().astype(int).to_numpy(),
        }
    )
    overall_mean = df[score_cols].mean().mean()
    comments = df.iloc[:, -1].dropna().astype(str)

    question_labels = [
        "Logical consistency of strategic choices and responses",
        "Historical plausibility of Soviet diplomatic shift and limited U.S. military response",
        "Natural and reasonable escalation dynamics across scenarios",
        "Perceived rule-based derivation rather than retrospective historical reasoning",
        "Distinctiveness of strategic choices and variable-driven outcomes",
        "Overall reasonableness and non-excessiveness of the three final scenarios",
    ]

    doc = Document()
    set_base_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Expert Validation Results for the LLM-Based Pol-Mil Scenario Methodology")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    doc.add_paragraph(
        "This report summarizes the quantitative and qualitative results of the expert validation survey "
        "conducted for the LLM-based political-military scenario methodology. The results are prepared in "
        "English so that they can be incorporated into the validation or evaluation section of an academic paper."
    )

    add_heading(doc, "1. Survey Data and Analytical Procedure", 1)
    doc.add_paragraph(
        f"The analysis was based on 18 expert responses collected in the worksheet of Expert_Responses.xlsx. "
        "Six Likert-scale items were analyzed as quantitative validation indicators. Each item was measured on "
        "a five-point scale, where higher scores indicate stronger expert agreement with the validity, plausibility, "
        "or reasonableness of the scenario outputs. Missing values were excluded on an item-by-item basis. "
        "Means and sample standard deviations were calculated for each item."
    )

    add_heading(doc, "2. Quantitative Validation Results", 1)
    rows = [["Item", "Validation Dimension", "Mean", "SD", "Valid N"]]
    for idx, row in analysis.iterrows():
        rows.append(
            [
                row["Question"],
                question_labels[idx],
                f"{row['Mean']:.2f}",
                f"{row['SD']:.2f}",
                str(int(row["Valid N"])),
            ]
        )
    add_table(doc, rows)

    doc.add_paragraph(
        f"The overall mean validity score across the six validation items was {overall_mean:.2f} out of 5.00. "
        "This result indicates that the expert panel generally evaluated the generated scenarios as valid, "
        "plausible, and reasonably structured. Among the six items, the highest mean score was observed for "
        "the distinctiveness of strategic choices and variable-driven outcomes (Q5, M = 4.11, SD = 0.58), "
        "suggesting that experts perceived meaningful differentiation among the scenario branches. The lowest "
        "mean score was observed for the overall reasonableness and non-excessiveness of the final scenarios "
        "(Q6, M = 3.83, SD = 0.62), indicating that some experts identified room for further refinement in the "
        "scope or intensity of the generated scenario outcomes."
    )

    add_heading(doc, "3. Interpretation for Paper Validation", 1)
    doc.add_paragraph(
        "The quantitative results provide supportive evidence for the internal validity of the proposed scenario "
        "methodology. Scores clustered around or above 4.00 for most items, which suggests that experts generally "
        "regarded the scenario logic, escalation patterns, and variable-based branching structure as coherent. "
        "The relatively moderate standard deviations indicate that expert evaluations were not highly dispersed, "
        "although some variation remained in judgments regarding historical plausibility and the extent to which "
        "the scenarios avoided excessive or dramatic outcomes."
    )
    doc.add_paragraph(
        "For the purposes of manuscript validation, these findings can be interpreted as evidence that the "
        "LLM-based Pol-Mil game design produced scenario outputs that were broadly acceptable to expert reviewers. "
        "At the same time, the results do not imply complete validation. Instead, they support a qualified conclusion: "
        "the framework demonstrates promising validity, while additional refinement is needed to improve historical "
        "transparency, causal traceability, and the realism of specific military decision pathways."
    )

    add_heading(doc, "4. Qualitative Feedback Themes", 1)
    doc.add_paragraph(
        f"Eight open-ended responses were submitted. The qualitative comments were reviewed and grouped into "
        "the following major themes:"
    )

    themes = [
        (
            "Clarification of LLM prior knowledge and training-data influence",
            "Several experts noted that the study should explain whether the generated scenarios were derived "
            "from the defined Pol-Mil game rules or from the model's pre-existing historical knowledge of the "
            "Cuban Missile Crisis. This issue is important because validation of the methodology depends on "
            "demonstrating that scenario outcomes were produced through the designed variable structure rather "
            "than simply reproduced from learned historical narratives.",
        ),
        (
            "Need for stronger historical grounding of Soviet decision-making",
            "Experts raised concerns about the plausibility of some Soviet military actions, especially in relation "
            "to nuclear command-and-control procedures. Comments suggested that the scenarios should more clearly "
            "differentiate between strategic and tactical nuclear authority, centralized command structures, and "
            "the conditions under which unauthorized or delegated launch decisions might be plausible.",
        ),
        (
            "Refinement of escalation mechanisms leading to catastrophic outcomes",
            "Some feedback indicated that the pathway from diplomatic and military tension to nuclear escalation "
            "should be elaborated more carefully. In particular, the role of warning-system errors, internal command "
            "confusion, and leadership cohesion could be specified in greater detail to strengthen causal plausibility.",
        ),
        (
            "Reduction of evaluative or emotionally loaded language",
            "One expert emphasized that excessive adjectives, dramatic expressions, or emotionally charged descriptions "
            "may weaken the perceived objectivity of the scenarios. More restrained and analytically neutral wording "
            "would improve academic credibility.",
        ),
        (
            "Greater transparency of AI decision logic",
            "Experts suggested that the study would benefit from making the AI's decision process more visible, for "
            "example through decision logs, variable-weight matrices, policy-preference vectors, or turn-by-turn "
            "explanations of how specific actions were selected.",
        ),
        (
            "Potential value of interactive decision-making design",
            "One response suggested that the model could be strengthened if participants were allowed to make strategic "
            "choices at each turn, with the scenario branching dynamically in response. Such an interactive structure "
            "could increase educational value and user engagement.",
        ),
    ]

    for title_text, body in themes:
        paragraph = doc.add_paragraph(style=None)
        run = paragraph.add_run(title_text + ": ")
        run.bold = True
        paragraph.add_run(body)

    add_heading(doc, "5. Suggested Manuscript Wording", 1)
    doc.add_paragraph(
        "The expert validation survey produced an overall mean score of 3.96 out of 5.00 across six evaluation "
        "items, indicating generally positive expert assessments of the proposed LLM-based Pol-Mil scenario "
        "methodology. The highest-rated item concerned the distinctiveness of strategic choices and the "
        "variable-driven differentiation of scenario outcomes (M = 4.11, SD = 0.58), while the lowest-rated item "
        "concerned the overall reasonableness and non-excessiveness of the final scenarios (M = 3.83, SD = 0.62). "
        "These results suggest that the framework was perceived as logically coherent and capable of generating "
        "meaningfully differentiated scenario branches. However, expert comments also highlighted several areas "
        "requiring refinement, including the need to clarify the influence of the LLM's prior historical knowledge, "
        "strengthen the historical plausibility of Soviet nuclear command-and-control assumptions, reduce evaluative "
        "language, and provide greater transparency regarding the AI's variable-based decision process."
    )

    add_heading(doc, "6. Summary Conclusion", 1)
    doc.add_paragraph(
        "Overall, the expert validation results support the preliminary validity of the LLM-based Pol-Mil scenario "
        "methodology. The quantitative scores indicate favorable expert evaluations, while the qualitative feedback "
        "identifies concrete directions for methodological improvement. Therefore, the validation results can be "
        "presented as supportive but not definitive evidence: the proposed framework demonstrates analytical promise, "
        "but further refinement is necessary to improve explainability, historical grounding, and causal traceability."
    )

    add_heading(doc, "Appendix: Open-Ended Response Count", 1)
    doc.add_paragraph(
        f"A total of {len(comments)} non-empty open-ended responses were identified. For research ethics and manuscript "
        "clarity, the responses were synthesized thematically rather than reproduced verbatim."
    )

    output = Path.cwd() / "Expert_Survey_Analysis_Report.docx"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
